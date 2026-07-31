"""
Visteon SAP CPI - Technical Specification Document Generator
==============================================================
Parses SAP Integration Suite (CPI) iFlow design-time ZIP exports and produces
a polished, corporate-styled Technical Specification .docx for each one -
cover page, auto-buildable table of contents, a redesigned process-flow
diagram, and consistently styled tables throughout.

Run:  python generate_ts_docs.py
Reads every *.zip in ZIP_DIRECTORY_PATH, writes one Technical Specifications_<iflow>.docx per ZIP
next to this script.
"""

import os
import glob
import re
import zipfile
import textwrap
import xml.etree.ElementTree as ET

import json
from google import genai

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# ===========================================================================
# LLM SETUP & EXTRACTION FUNCTION
# ===========================================================================
# Initialize Gemini Client (reads GEMINI_API_KEY from environment variables)
try:
    llm_client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
except Exception as e:
    llm_client = None
    print(f"Warning: Gemini client not initialized. {e}")

def extract_adapters_and_systems_with_llm(bpmn_xml_content, iflow_name=""):
    """
    Passes raw BPMN iFlow XML to Gemini to accurately identify Sender/Receiver systems
    and their exact Adapter Types.
    """
    if not llm_client:
        return None

    prompt = f"""
    You are an expert SAP CPI / Integration Suite Architect.
    Analyze the following raw BPMN iFlow XML content for the iFlow named "{iflow_name}".

    Your task is to identify and return a JSON object with the following keys:
    1. "sender_system": The source system triggering or sending data.
    2. "sender_adapter_type": The exact adapter on the inbound connection.
    3. "receiver_system": The primary target business system.
    4. "receiver_adapter_type": The exact adapter/protocol on the main outbound connection.
    5. "execution_mode": If the XML contains <timerStartEvent>, this MUST be "Scheduled / Batched". Otherwise, "Real-time / Event-driven".
    6. "frequency": "Configured Cron Schedule" if Scheduled, else "Immediate upon request".
    7. "business_overview": A concise business-focused summary of the integration's purpose.
    8. "interface_description": A concise technical description of how the interface works.
    9. "processing_logic": An ARRAY OF STRINGS, where each string is a highly detailed step (e.g., "1. Polling: ...", "2. Mapping: ..."). Explicitly include technical nuances like dynamic filename generation, Data Store operations, and Archiving logic. Keep one distinct point per array item.
    10. "diagram_nodes": An array of objects representing the MAIN pipeline flow in exact execution order. Each object must have:
        - "name": The label of the node (e.g., "QAD", "Transform Payload", "Fifth Third Bank", "Archive Data", "End").
        - "subtype": The top banner text (e.g., "SFTP Sender", "Processing Step", "SFTP Receiver", "End Event").
        - "node_type": Strictly one of: "Sender", "Receiver", "Processing", "End".
        *CRITICAL*: Place the Receiver exactly where it happens in the flow (e.g., step 3). Do NOT force it to the end if archiving or emails happen after it. End the array with an "End" node.
    11. "exception_nodes": An array of objects representing the LOCAL EXCEPTION SUBPROCESS flow in exact execution order. Each object must have:
        - "name": The exact step name.
        - "subtype": The top banner text (e.g., "Exception Handling").

    BPMN XML Content:
    {bpmn_xml_content[:15000]}
    """

    try:
        response = llm_client.models.generate_content(
            model='gemini-3.5-flash-lite',
            contents=prompt,
        )

        raw_text = getattr(response, "text", "") or ""
        if not raw_text.strip():
            return None

        # Safely extract the JSON block directly from the raw response
        # This completely avoids the newline splitting error.
        match = re.search(r"\{.*\}", raw_text, re.DOTALL)
        if match:
            return json.loads(match.group(0))

        return None

    except Exception as e:
        print(f"LLM Adapter Extraction warning: {e}")
        return None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PRIMARY_LOGO_PATH = os.path.join(BASE_DIR, "Visteon_Logo.jpeg")
ZIP_DIRECTORY_PATH = r"/Volumes/Drive D/CPI Python Script"

# ===========================================================================
# THEME - one place to tune the whole document's look
# ===========================================================================
NAVY       = "1F3864"   # headings, primary structure
NAVY_SOFT  = "2E5395"   # secondary accents / diagram nodes
ORANGE     = "F2661D"   # Visteon accent - section numbers, rules, sub-titles
SLATE      = "44546A"   # body label text
GREEN      = "2E7D32"   # sender / receiver endpoints
RED        = "C0392B"   # exception / error path
PURPLE     = "7A4FC0"   # data mapping nodes
AMBER      = "B08900"   # converter / store nodes
BG_LIGHT   = "F4F7FB"   # zebra row / panel fill
BG_ERR     = "FDF1EF"   # exception panel fill
GRID       = "D7DEE8"   # hairline borders
WHITE      = "FFFFFF"
TEXT_BODY  = "333333"

FONT = "Calibri"

TYPE_COLORS = {
    "Groovy Script":     ORANGE,
    "Data Mapping":      PURPLE,
    "Lookup / VM":       "0E7C86",
    "Converter / Store": AMBER,
    "Processing Step":   NAVY_SOFT,
    "Exception Handler": RED,
}


def rgb(hexstr):
    return RGBColor.from_string(hexstr)


def resolve_asset_path(path):
    if not path:
        return None
    if os.path.isabs(path):
        return path
    candidates = [path, os.path.join(BASE_DIR, path), os.path.join(os.getcwd(), path)]
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return os.path.join(BASE_DIR, path)


def _is_timer_start_element(elem):
    tag_clean = elem.tag.split('}')[-1]
    if tag_clean == 'timerStartEvent':
        return True
    if tag_clean == 'startEvent':
        for child in elem:
            if child.tag.split('}')[-1] == 'timerEventDefinition':
                return True
    return False


def sanitize_output_name(name):
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", name or "iflow").strip("._-")
    return safe_name or "iflow"


RECEIVER_ONLY_ADAPTERS = {
    "idoc", "odata", "rfc", "jdbc", "sap", "s4", "s4hana", "edi", "as2", "jms", "mq",
    "api", "webservice", "rest"
}

SENDER_ONLY_ADAPTERS = {"timer"}
SENDER_PREFERRED_ADAPTERS = {"sftp", "ftp", "file", "http", "https", "soap", "rest", "api", "webservice", "mail", "smtp", "imap", "pop"}

# Adapters like SFTP/FTP/File/HTTP/Mail can legitimately act as EITHER the
# sender (polling/reading a source) or the receiver (writing/delivering an
# output) depending on the flow - they should never be assumed to always be
# "sender" channels. Keyword hints from the channel/step naming are used to
# disambiguate; when there is no hint, the fallback in infer_channel_role()
# below takes the presence of a timer start event into account.
SENDER_INTENT_KEYWORDS = [
    "sender", "source", "inbound", "start", "trigger", "schedule", "sched",
    "poll", "receive", "listen", "read", "get", "download", "pickup",
    "pick up", "fetch", "consume", "request"
]
RECEIVER_INTENT_KEYWORDS = [
    "receiver", "target", "outbound", "response", "destination", "send",
    "deliver", "write", "put", "upload", "push", "post", "create", "update"
]


def normalize_adapter_type(adapter_type):
    if not adapter_type:
        return "Adapter"
    normalized = adapter_type.strip().lower()
    if "idoc" in normalized:
        return "IDoc"
    if "odata" in normalized:
        return "OData"
    if "https" in normalized:
        return "HTTPS"
    if "http" in normalized:
        return "HTTP"
    if "sftp" in normalized:
        return "SFTP"
    if "ftp" in normalized:
        return "FTP"
    if "file" in normalized:
        return "File"
    if "soap" in normalized:
        return "SOAP"
    if "rest" in normalized or "api" in normalized or "webservice" in normalized:
        return "REST"
    if "mail" in normalized or "smtp" in normalized or "pop" in normalized or "imap" in normalized:
        return "Mail"
    if "timer" in normalized or "schedule" in normalized or "cron" in normalized:
        return "Timer"
    if "jdbc" in normalized:
        return "JDBC"
    if "rfc" in normalized:
        return "RFC"
    if "as2" in normalized:
        return "AS2"
    if "jms" in normalized:
        return "JMS"
    if "edi" in normalized:
        return "EDI"
    return adapter_type


def _normalize_text(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())


def _looks_like_placeholder(value):
    normalized = _normalize_text(value).lower()
    if not normalized:
        return True
    normalized = re.sub(r"[^a-z0-9]+", "", normalized)
    return normalized in {"test", "sample", "dummy", "temp", "tbd", "unknown", "na", "n/a", "sender", "receiver", "sendersystem", "receiversystem"}


def build_processing_logic_from_pipeline(iflow_data):
    iflow_data = iflow_data or {}
    sender = _normalize_text(iflow_data.get("sender_system") or "sender")
    receiver = _normalize_text(iflow_data.get("receiver_system") or "receiver")

    pipeline_nodes = list(iflow_data.get("main_pipeline_nodes") or [])
    if pipeline_nodes:
        step_names = []
        for name, step_type in pipeline_nodes[:6]:
            cleaned = _normalize_text(name)
            if cleaned:
                step_names.append(cleaned)
        if step_names:
            return (
                f"The flow receives data from {sender}, executes the processing steps {', '.join(step_names)} "
                f"in SAP CPI, and delivers the result to {receiver}."
            )

    return (
        f"The flow receives data from {sender} and processes it in SAP CPI before delivering the result to {receiver}."
    )


def _derive_diagram_steps_from_logic(processing_logic, iflow_data=None):
    text = _normalize_text(processing_logic or "")
    if not text:
        text = build_processing_logic_from_pipeline(iflow_data or {})

    lowered = text.lower()
    steps = []

    if any(term in lowered for term in ["validate", "validation", "check", "verify"]):
        steps.append("Validate the inbound payload and required fields")
    if any(term in lowered for term in ["map", "mapping", "transform", "transformation", "format"]):
        steps.append("Map and transform the payload to the target structure")
    if any(term in lowered for term in ["enrich", "lookup", "reference", "standard", "context"]):
        steps.append("Enrich the payload with reference or business data")
    if any(term in lowered for term in ["filter", "route", "split", "condition", "branch"]):
        steps.append("Filter or route the message based on business rules")
    if any(term in lowered for term in ["call", "invoke", "send", "post", "request", "endpoint", "target", "deliver"]):
        steps.append("Invoke the target endpoint and deliver the payload")

    if not steps:
        steps.append("Receive the source message")
        steps.append("Process the payload in SAP CPI")
        steps.append("Deliver the processed result to the target")

    return steps[:6]


def apply_llm_interface_intelligence(iflow_data, llm_data=None):
    merged = dict(iflow_data or {})
    llm_data = llm_data or {}

    # 1. Text Normalization for Overviews
    business_overview = _normalize_text(llm_data.get("business_overview") or llm_data.get("overview") or merged.get("business_overview"))
    interface_description = _normalize_text(llm_data.get("interface_description") or llm_data.get("description") or merged.get("interface_description"))
    
    if business_overview:
        merged["business_overview"] = business_overview
    if interface_description:
        merged["interface_description"] = interface_description

    # 2. Preserve Array Structure for Processing Logic
    p_logic = llm_data.get("processing_logic") or llm_data.get("high_level_processing_logic") or merged.get("processing_logic")
    if isinstance(p_logic, list):
        # Keep as list for bullet-point formatting
        merged["processing_logic"] = [str(step).strip() for step in p_logic]
    elif p_logic:
        # Fallback to string if LLM failed to return a list
        merged["processing_logic"] = _normalize_text(p_logic)

    # 3. Pass through structured JSON arrays for the diagram
    if isinstance(llm_data.get("diagram_nodes"), list):
        merged["diagram_nodes"] = llm_data.get("diagram_nodes")
        
    if isinstance(llm_data.get("exception_nodes"), list):
        merged["exception_nodes"] = llm_data.get("exception_nodes")

    # 4. Standard Metadata Overrides
    if llm_data.get("sender_system"):
        merged["sender_system"] = _normalize_text(llm_data.get("sender_system"))
    if llm_data.get("receiver_system"):
        merged["receiver_system"] = _normalize_text(llm_data.get("receiver_system"))
    if llm_data.get("sender_adapter_type"):
        merged["sender_adapter_type"] = _normalize_text(llm_data.get("sender_adapter_type"))
    if llm_data.get("receiver_adapter_type"):
        merged["receiver_adapter_type"] = _normalize_text(llm_data.get("receiver_adapter_type"))
    if llm_data.get("execution_mode"):
        merged["execution_mode"] = _normalize_text(llm_data.get("execution_mode"))
    if llm_data.get("frequency"):
        merged["frequency"] = _normalize_text(llm_data.get("frequency"))

    return merged


def apply_user_metadata(iflow_data, user_inputs=None):
    merged = dict(iflow_data or {})
    user_inputs = user_inputs or {}

    prepared_by = _normalize_text(user_inputs.get("prepared_by") or user_inputs.get("preparedBy") or merged.get("prepared_by")) or "Not provided"
    reviewed_by = _normalize_text(user_inputs.get("reviewed_by") or user_inputs.get("reviewedBy") or merged.get("reviewed_by")) or "Not provided"
    approved_by = _normalize_text(user_inputs.get("approved_by") or user_inputs.get("approvedBy") or merged.get("approved_by")) or "Not provided"
    effective_date = _normalize_text(user_inputs.get("effective_date") or user_inputs.get("effectiveDate") or merged.get("effective_date")) or "TBD"
    description = _normalize_text(user_inputs.get("description") or merged.get("description")) or "No description provided."

    direction = _normalize_text(user_inputs.get("direction") or merged.get("direction"))
    if direction.lower() in {"inbound", "outbound"}:
        direction = direction.title()
    else:
        direction = merged.get("direction") or "Inbound"
        if not isinstance(direction, str):
            direction = "Inbound"
        direction = direction.title()

    sync_async = _normalize_text(user_inputs.get("sync_async") or user_inputs.get("syncAsync") or merged.get("synchronous_asynchronous"))
    if sync_async.lower() in {"sync", "synchronous"}:
        sync_async = "Synchronous"
    elif sync_async.lower() in {"async", "asynchronous"}:
        sync_async = "Asynchronous"
    else:
        sync_async = merged.get("synchronous_asynchronous") or "Synchronous"

    sender_system = _normalize_text(
        user_inputs.get("source_system") or user_inputs.get("sourceSystem") or merged.get("sender_system")
    )
    receiver_system = _normalize_text(
        user_inputs.get("target_system") or user_inputs.get("targetSystem") or merged.get("receiver_system")
    )
    if _looks_like_placeholder(sender_system):
        sender_system = "Source System"
    if _looks_like_placeholder(receiver_system):
        receiver_system = "Target System"

    sender_adapter = _normalize_text(merged.get("sender_adapter_type"))
    receiver_adapter = _normalize_text(merged.get("receiver_adapter_type"))
    if _looks_like_placeholder(sender_adapter):
        sender_adapter = "HTTP/HTTPS"
    if _looks_like_placeholder(receiver_adapter):
        receiver_adapter = "REST / OData"

    package_name = _normalize_text(user_inputs.get("package_name") or merged.get("package_name")) or "SAP Integration Package"
    doc_version = _normalize_text(user_inputs.get("doc_version")) or "1.0"

    merged.update({
        "package_name": package_name,
        "doc_version": doc_version,
        "prepared_by": prepared_by,
        "reviewed_by": reviewed_by,
        "approved_by": approved_by,
        "effective_date": effective_date,
        "description": description,
        "direction": direction,
        "synchronous_asynchronous": sync_async,
        "sender_system": sender_system,
        "receiver_system": receiver_system,
        "sender_adapter_type": sender_adapter,
        "receiver_adapter_type": receiver_adapter,
    })

    # Respect whatever execution mode the parser already determined (e.g.
    # "Scheduled / Batched" for a timer-triggered iFlow, "Real-time /
    # Event-driven" otherwise) instead of forcing it to Real-time whenever
    # the direction happens to be Outbound - a scheduler-driven Outbound
    # interface (e.g. a nightly file push) is still Scheduled/Batched, not
    # Real-time.
    merged["execution_mode"] = merged.get("execution_mode") or "Real-time / Event-driven"

    if not merged.get("frequency"):
        merged["frequency"] = "On demand / operational schedule"

    return merged


def infer_channel_role(channel_name, channel_id, source_ref, target_ref, adapter_type, is_timer_iflow):
    text = f"{channel_name} {channel_id} {source_ref} {target_ref}".lower()
    adapter_lower = adapter_type.strip().lower()

    # Adapters that are unambiguous by nature.
    if adapter_lower in RECEIVER_ONLY_ADAPTERS:
        return "receiver"
    if adapter_lower in SENDER_ONLY_ADAPTERS:
        return "sender"

    # Explicit source/target refs from the BPMN participants win next.
    if any(term in target_ref.lower() for term in ["target", "receiver", "endpoint", "destination"]):
        return "receiver"
    if any(term in source_ref.lower() for term in ["source", "sender", "start", "trigger"]):
        return "sender"

    # Adapters such as SFTP/FTP/File/HTTP/Mail can be sender OR receiver
    # channels depending on the flow - e.g. an SFTP channel that WRITES an
    # outbound file is a receiver, not a sender, even though SFTP can also be
    # used to poll an inbound file. Use naming keywords to disambiguate
    # instead of assuming these adapters are always senders.
    if adapter_lower in SENDER_PREFERRED_ADAPTERS:
        if any(term in text for term in RECEIVER_INTENT_KEYWORDS):
            return "receiver"
        if any(term in text for term in SENDER_INTENT_KEYWORDS):
            return "sender"
        if is_timer_iflow:
            # The timer start event is already the sender/trigger for this
            # flow, so an unlabeled downstream channel is far more likely to
            # be where the processed data is delivered to.
            return "receiver"
        return "sender"

    if any(term in text for term in RECEIVER_INTENT_KEYWORDS) or any(
        term in text for term in ["idoc", "odata", "sap", "s4", "s4hana", "db", "erp", "rfc", "jdbc", "jms", "mq", "as2", "edi", "rest", "api", "webservice"]
    ):
        return "receiver"
    if any(term in text for term in SENDER_INTENT_KEYWORDS):
        return "sender"

    return None


def _security_material_label(key, value):
    if not key:
        return "Security Material"
    label = key.replace('_', ' ').replace('-', ' ').title()
    if label.lower() in ["value", "string", "text"]:
        return "Security Material"
    return label


def _add_security_material(data, key, value):
    if not value:
        return
    if isinstance(value, str):
        value = value.strip()
    if not value:
        return
    value = str(value)
    # Avoid repeated duplicates
    if any(value == existing_value for _, existing_value in data["security_materials"]):
        return
    material_type = _security_material_label(key, value)
    data["security_materials"].append((material_type, value))


# ===========================================================================
# CPI ZIP & BPMN PARSER  (unchanged logic - already handles arbitrary iFlow
# exports dynamically; only lightly cleaned up)
# ===========================================================================

import os
import glob
import zipfile
import textwrap
import xml.etree.ElementTree as ET

def parse_cpi_iflow_zip(zip_file_path, user_inputs=None):
    base_filename = os.path.splitext(os.path.basename(zip_file_path))[0]

    data = {
        "iflow_name": base_filename,
        "package_name": "SAP Integration Package",
        "sender_system": "Sender System",
        "receiver_system": "Receiver System",
        "sender_adapter_type": "HTTPS",
        "receiver_adapter_type": "OData",
        "execution_mode": "Real-time / Event-driven",
        "synchronous_asynchronous": "Synchronous",
        "frequency": "Immediate upon request",
        "direction": "Inbound",
        "prepared_by": "",
        "reviewed_by": "",
        "approved_by": "",
        "effective_date": "",
        "description": "",
        "main_pipeline_nodes": [],
        "exception_nodes": [],
        "groovy_scripts": [],
        "xslt_scripts": [],
        "mappings": [],
        "properties": {},
        "sender_params": [],
        "receiver_params": [],
        "security_materials": []
    }

    try:
        with zipfile.ZipFile(zip_file_path, 'r') as z:
            file_list = z.namelist()

            # 1. Parse metainfo.prop
            metainfo_file = [f for f in file_list if f.endswith('metainfo.prop')]
            if metainfo_file:
                content = z.read(metainfo_file[0]).decode('utf-8', errors='ignore')
                for line in content.splitlines():
                    if '=' in line and not line.strip().startswith('#'):
                        k, v = line.split('=', 1)
                        k, v = k.strip(), v.strip()
                        if k in ["SymbolicName", "Bundle-Name", "Name"] and v:
                            data["iflow_name"] = v
                        elif k in ["OriginBundle-SymbolicName", "PackageName"] and v:
                            data["package_name"] = v

            # 2. Extract script & mapping files
            for f in file_list:
                fname = os.path.basename(f)
                if not fname:
                    continue
                if f.endswith(".groovy"):
                    data["groovy_scripts"].append(fname)
                elif f.endswith(".xsl") or f.endswith(".xslt"):
                    data["xslt_scripts"].append(fname)
                elif f.endswith(".mmap") or f.endswith(".mmap.xml") or f.endswith(".map"):
                    data["mappings"].append(fname)

            # 3. Parse parameters.prop
            param_files = [f for f in file_list if f.endswith('parameters.prop')]
            if param_files:
                prop_content = z.read(param_files[0]).decode('utf-8', errors='ignore')
                for line in prop_content.splitlines():
                    if '=' in line and not line.strip().startswith('#'):
                        k, v = line.split('=', 1)
                        data["properties"][k.strip()] = v.strip()
                        if any(term in k.lower() for term in ["credential", "alias", "user", "key", "auth", "password", "certificate", "material"]):
                            _add_security_material(data, k.strip(), v.strip())

            # 4. Parse BPMN iFlow XML (*.iflw / *.ifbw / component.xml)
            bpmn_files = [f for f in file_list if f.endswith('.ifbw') or f.endswith('.iflw') or f.endswith('component.xml')]
            bpmn_xml_str = None
            if bpmn_files:
                bpmn_xml_bytes = z.read(bpmn_files[0])
                bpmn_xml_str = bpmn_xml_bytes.decode('utf-8', errors='ignore')

                tree = ET.fromstring(bpmn_xml_bytes)

                participants = {}
                is_timer_iflow = False
                receiver_candidates = []
                sender_candidates = []
                channel_candidates = []

                # First pass: Check for Timer Start Events
                for elem in tree.iter():
                    if _is_timer_start_element(elem):
                        is_timer_iflow = True
                        break

                # Set initial defaults based on start event
                if is_timer_iflow:
                    data["sender_system"] = "SAP CPI Scheduler"
                    data["sender_adapter_type"] = "Timer"
                    data["execution_mode"] = "Scheduled / Batched"
                    data["synchronous_asynchronous"] = "Asynchronous"
                    data["frequency"] = "Configured Cron Schedule"
                    # Direction (Inbound/Outbound) is determined once all
                    # channels have been parsed - see below - instead of
                    # being hard-coded to a generic "Scheduled" label.

                # Parse participants and channels dynamically
                for elem in tree.iter():
                    tag_clean = elem.tag.split('}')[-1]

                    if tag_clean == 'participant':
                        p_name = elem.attrib.get('name', '').strip()
                        p_id = elem.attrib.get('id', '').strip()
                        if p_name and p_name not in ["Integration Process", "Participant", "Enclosing Process"]:
                            participants[p_id] = p_name

                    elif tag_clean == 'channel':
                        adapter_type = normalize_adapter_type(elem.attrib.get('type', 'Adapter'))
                        channel_name = elem.attrib.get('name', '')
                        channel_id = elem.attrib.get('id', '')
                        source_ref = elem.attrib.get('sourceRef', '')
                        target_ref = elem.attrib.get('targetRef', '')

                        channel_props = []
                        for prop in elem.iter():
                            p_clean = prop.tag.split('}')[-1]
                            if p_clean in ['property', 'attribute']:
                                k = prop.attrib.get('key') or prop.findtext('./key')
                                v = prop.attrib.get('value') or prop.findtext('./value')
                                if k and v:
                                    channel_props.append((k, v))
                                    if any(term in k.lower() for term in ["credential", "alias", "user", "key", "auth", "password", "certificate", "material"]):
                                        _add_security_material(data, k, v)
                            elif p_clean in ['key', 'value'] and elem.tag.split('}')[-1] == 'property':
                                # support nested property elements from some CPI XML schemas
                                k = elem.findtext('./key') or prop.attrib.get('key')
                                v = elem.findtext('./value') or prop.attrib.get('value')
                                if k and v:
                                    channel_props.append((k, v))
                                    if any(term in k.lower() for term in ["credential", "alias", "user", "key", "auth", "password", "certificate", "material"]):
                                        _add_security_material(data, k, v)

                        channel_candidates.append({
                            "adapter_type": adapter_type,
                            "channel_name": channel_name,
                            "channel_id": channel_id,
                            "source_ref": source_ref,
                            "target_ref": target_ref,
                            "channel_props": channel_props,
                            "channel_role": infer_channel_role(channel_name, channel_id, source_ref, target_ref, adapter_type, is_timer_iflow)
                        })

                        if channel_candidates[-1]["channel_role"] == 'sender':
                            if is_timer_iflow:
                                # The timer start event is already the
                                # sender/trigger - don't let a downstream
                                # channel rename the sender adapter away from
                                # "Timer". Still capture its parameters if
                                # nothing has been recorded yet.
                                if not data["sender_params"]:
                                    data["sender_params"] = channel_props
                            else:
                                data["sender_params"] = channel_props
                                data["sender_adapter_type"] = adapter_type
                                data["execution_mode"] = "Real-time / Event-driven"
                        elif channel_candidates[-1]["channel_role"] == 'receiver':
                            data["receiver_params"].extend(channel_props)
                            data["receiver_adapter_type"] = adapter_type
                        else:
                            # Fall back to sensible defaults when channel role is ambiguous
                            if any(term in adapter_type.lower() for term in ["idoc", "odata", "rfc", "jdbc", "sap", "s4", "s4hana", "edi", "as2", "jms", "mq"]):
                                data["receiver_params"].extend(channel_props)
                                data["receiver_adapter_type"] = adapter_type
                            elif not is_timer_iflow and adapter_type.lower() == "sftp":
                                data["sender_params"] = channel_props
                                data["sender_adapter_type"] = adapter_type
                                data["execution_mode"] = "Real-time / Event-driven"
                            elif not is_timer_iflow and adapter_type.lower() in SENDER_PREFERRED_ADAPTERS:
                                data["sender_params"] = channel_props
                                data["sender_adapter_type"] = adapter_type
                                data["execution_mode"] = "Real-time / Event-driven"
                            else:
                                data["receiver_params"].extend(channel_props)
                                data["receiver_adapter_type"] = adapter_type

                # Dynamic System Name Resolution based on Participant roles / Names
                for p_id, p_name in participants.items():
                    p_name_lower = p_name.lower()
                    p_id_lower = p_id.lower()

                    if any(s in p_id_lower or s in p_name_lower for s in ["sender", "source", "cimple", "mes", "client", "http"]):
                        sender_candidates.append(p_name)
                    elif any(r in p_id_lower or r in p_name_lower for r in ["receiver", "target", "sap", "s4", "s4hana", "gbc", "odata", "db"]):
                        receiver_candidates.append(p_name)
                    else:
                        # Fallback heuristic
                        receiver_candidates.append(p_name)

                if sender_candidates and not is_timer_iflow:
                    data["sender_system"] = sender_candidates[0]
                elif not is_timer_iflow:
                    # Look inside iflow name for source clues (e.g. CELOAD_CIMPLE_To_S4HANA -> CIMPLE)
                    if "cimple" in base_filename.lower():
                        data["sender_system"] = "CIMPLE"

                if receiver_candidates:
                    # Remove duplicates while preserving order
                    unique_receivers = list(dict.fromkeys(receiver_candidates))
                    data["receiver_system"] = " / ".join(unique_receivers)

                # Second pass: if no clear sender/receiver was found, use channel candidate heuristics
                if (not data["sender_params"] or not data["receiver_params"]) and channel_candidates:
                    for candidate in channel_candidates:
                        if candidate["channel_role"] == "sender" and not is_timer_iflow:
                            data["sender_params"] = candidate["channel_props"]
                            data["sender_adapter_type"] = candidate["adapter_type"]
                        elif candidate["channel_role"] == "receiver" and not data["receiver_params"]:
                            data["receiver_params"] = candidate["channel_props"]
                            data["receiver_adapter_type"] = candidate["adapter_type"]

                # If still ambiguous, prioritize explicit adapter names
                if not data["receiver_params"]:
                    for candidate in channel_candidates:
                        if candidate["adapter_type"].lower() == "idoc":
                            data["receiver_params"] = candidate["channel_props"]
                            data["receiver_adapter_type"] = candidate["adapter_type"]
                            break
                # A bare SFTP fallback for "sender" only makes sense outside
                # timer flows - for a scheduler-driven flow the timer is
                # already the sender, and an SFTP channel here is almost
                # always where the output is being delivered to.
                if not data["sender_params"] and not is_timer_iflow:
                    for candidate in channel_candidates:
                        if candidate["adapter_type"].lower() == "sftp":
                            data["sender_params"] = candidate["channel_props"]
                            data["sender_adapter_type"] = candidate["adapter_type"]
                            break

                # For timer/scheduler-driven flows, determine the true
                # interface direction from where the processed data actually
                # lands, instead of a generic "Scheduled" label - this keeps
                # Inbound/Outbound meaningful even for scheduler interfaces.
                if is_timer_iflow:
                    receiver_adapter_lower = (data.get("receiver_adapter_type") or "").lower()
                    sap_bound_indicators = ["idoc", "odata", "rfc", "jdbc", "sap", "s4", "s4hana"]
                    if any(term in receiver_adapter_lower for term in sap_bound_indicators):
                        data["direction"] = "Inbound"
                    else:
                        data["direction"] = "Outbound"

                # Extract steps logic
                extracted_steps = []
                exception_steps = []

                for elem in tree.iter():
                    tag_clean = elem.tag.split('}')[-1]
                    name = elem.attrib.get('name', '').strip()

                    if tag_clean in ['callActivity', 'serviceTask', 'scriptTask', 'sendTask', 'subProcess', 'task']:
                        if not name or name in ["Start", "End", "Start 1", "End 1", "Start 2", "End 2"]:
                            continue

                        if any(e in name.lower() or e in elem.attrib.get('id', '').lower() for e in ['exception', 'error', 'fault']):
                            exception_steps.append((name, "Exception Handler"))
                        else:
                            if any(w in name.lower() for w in ["groovy", "script", "gs_"]):
                                extracted_steps.append((name, "Groovy Script"))
                            elif any(w in name.lower() for w in ["xslt", "mapping", "mm_", "map"]):
                                extracted_steps.append((name, "Data Mapping"))
                            elif any(w in name.lower() for w in ["lookup", "vm", "value"]):
                                extracted_steps.append((name, "Lookup / VM"))
                            elif any(w in name.lower() for w in ["datastore", "payload", "json", "xml", "convert", "jdbc"]):
                                extracted_steps.append((name, "Converter / Store"))
                            else:
                                extracted_steps.append((name, "Processing Step"))

                data["main_pipeline_nodes"] = extracted_steps
                data["exception_nodes"] = exception_steps

    except Exception as e:
        print(f"WARNING: could not fully parse ZIP '{os.path.basename(zip_file_path)}': {e}")

    # Fallback overrides using explicit properties or naming conventions
    for k, v in data["properties"].items():
        if "receiver" in k.lower() or "target" in k.lower():
            data["receiver_system"] = v
        elif "sender" in k.lower() or "source" in k.lower():
            data["sender_system"] = v

    if not data.get("processing_logic"):
        data["processing_logic"] = build_processing_logic_from_pipeline(data)

    if not data.get("diagram_steps"):
        data["diagram_steps"] = _derive_diagram_steps_from_logic(data.get("processing_logic"), data)

    # =======================================================================
    # LLM OVERRIDE FOR ACCURATE ADAPTER & SYSTEM DETECTION
    # =======================================================================
    if bpmn_xml_str:
        llm_extracted = extract_adapters_and_systems_with_llm(bpmn_xml_str, data["iflow_name"])

        if llm_extracted:
            print(f"-> LLM successfully extracted adapters for: {data['iflow_name']}")
            data = apply_llm_interface_intelligence(data, llm_extracted)

    return apply_user_metadata(data, user_inputs)

# ===========================================================================
# PROCESS FLOW DIAGRAM - redesigned: card-style nodes with a coloured type
# banner, numbered step badges, a dedicated exception-flow swimlane and a
# legend, all rendered in the document's theme colours.
# ===========================================================================

def _wrap_label(text, width_chars, max_lines=6):
    """Wrap text to fit inside a fixed-width box without truncating with ellipses."""
    text = (text or "").strip().replace("_", " ")
    # Removed the placeholder="..." so text fully prints
    # break_long_words=False prevents weird hyphenation on technical terms
    lines = textwrap.wrap(text, width=width_chars, max_lines=max_lines, break_long_words=False)
    return lines or [""]


def _rounded_node(ax, x, y, w, h, label_lines, subtype_lines, color, number=None, zorder=5):
    ax.add_patch(patches.FancyBboxPatch(
        (x - w / 2 + 0.045, y - h / 2 - 0.045), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=0, facecolor="#00102A", alpha=0.14, zorder=zorder))
    ax.add_patch(patches.FancyBboxPatch(
        (x - w / 2, y - h / 2), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=1.1, edgecolor=f"#{color}", facecolor=f"#{WHITE}", zorder=zorder + 1))

    banner_h = 0.16 + 0.14 * len(subtype_lines)
    ax.add_patch(patches.FancyBboxPatch(
        (x - w / 2, y + h / 2 - banner_h), w, banner_h,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=0, facecolor=f"#{color}", zorder=zorder + 2))
    ax.add_patch(patches.Rectangle(
        (x - w / 2, y + h / 2 - banner_h), w, banner_h * 0.5,
        linewidth=0, facecolor=f"#{color}", zorder=zorder + 2))

    # subtype text (wrapped, centred in the banner)
    sub_y0 = y + h / 2 - banner_h / 2 + (len(subtype_lines) - 1) * 0.07
    for i, line in enumerate(subtype_lines):
        ax.text(x, sub_y0 - i * 0.14, line.upper(), fontsize=6.1, fontweight='bold',
                color=f"#{WHITE}", ha='center', va='center', zorder=zorder + 3,
                fontfamily='DejaVu Sans')

    # label text (wrapped, centred in the remaining body area)
    body_h = h - banner_h
    body_cy = y - banner_h / 2
    label_y0 = body_cy + (len(label_lines) - 1) * 0.085
    for i, line in enumerate(label_lines):
        ax.text(x, label_y0 - i * 0.17, line, fontsize=7.2, fontweight='bold',
                color=f"#{SLATE}", ha='center', va='center', zorder=zorder + 3,
                fontfamily='DejaVu Sans')

    if number is not None:
        ax.add_patch(patches.Circle((x - w / 2 + 0.02, y + h / 2 + 0.02), 0.16,
                     facecolor=f"#{NAVY}", edgecolor=f"#{WHITE}", linewidth=1.4, zorder=zorder + 4))
        ax.text(x - w / 2 + 0.02, y + h / 2 + 0.02, str(number), fontsize=7.5, fontweight='bold',
                color=f"#{WHITE}", ha='center', va='center', zorder=zorder + 5)


def _arrow(ax, x1, y1, x2, y2, color, lw=1.6):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=f"#{color}", lw=lw,
                                 mutation_scale=13, shrinkA=0, shrinkB=0), zorder=6)
def draw_enterprise_diagram(iflow_data, output_img_path="cpi_flow_diagram.png"):
    output_dir = os.path.dirname(output_img_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    diagram_nodes = iflow_data.get("diagram_nodes")
    
    # Fallback mapping if LLM failed to return structured diagram nodes
    if not diagram_nodes or not isinstance(diagram_nodes, list):
        diagram_nodes = [{"name": iflow_data.get("sender_system", "Sender"), "subtype": f"{iflow_data.get('sender_adapter_type', 'HTTP')} Sender", "node_type": "Sender"}]
        # Limit fallback processing steps to 3 to maintain conciseness
        for name, styp in iflow_data.get("main_pipeline_nodes", [])[:3]:
            diagram_nodes.append({"name": name, "subtype": styp, "node_type": "Processing"})
        diagram_nodes.append({"name": iflow_data.get("receiver_system", "Receiver"), "subtype": f"{iflow_data.get('receiver_adapter_type', 'HTTP')} Receiver", "node_type": "Receiver"})

    # ENFORCE CONCISENESS: Limit to max 6 nodes to keep diagram compact
    if len(diagram_nodes) > 6:
        diagram_nodes = diagram_nodes[:6]
        # Ensure the last node makes sense if we truncated
        if diagram_nodes[-1].get("node_type") not in ["Receiver", "End"]:
            diagram_nodes[-1] = {"name": "End Process", "subtype": "End Event", "node_type": "End"}

    total = len(diagram_nodes)
    
    # Compact sizing settings
    node_w = 2.1
    gap = node_w + 0.5 
    fig_w = max(12.5, total * gap + 1.5)
    fig_h = 6.4 
    
    LABEL_WRAP_CHARS = 18
    SUBTYPE_WRAP_CHARS = 18

    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=220)
    ax.set_xlim(0, fig_w)
    ax.set_ylim(0, fig_h)
    ax.axis('off')
    fig.patch.set_facecolor(f"#{WHITE}")

    iflow_name = iflow_data["iflow_name"]

    # Title bar
    ax.add_patch(patches.FancyBboxPatch((0.35, fig_h - 0.62), fig_w - 0.7, 0.5,
                 boxstyle="round,pad=0.02,rounding_size=0.08",
                 linewidth=0, facecolor=f"#{NAVY}", zorder=1))
    ax.text(0.6, fig_h - 0.37, f"Integration Flow  \u2014  {iflow_name}",
            fontsize=10.5, fontweight='bold', color=f"#{WHITE}", va='center', zorder=2,
            fontfamily='DejaVu Sans')

    # Main pipeline panel
    main_top, main_bottom = fig_h - 0.85, 2.05
    ax.add_patch(patches.FancyBboxPatch((0.35, main_bottom), fig_w - 0.7, main_top - main_bottom,
                 boxstyle="round,pad=0.02,rounding_size=0.06",
                 linewidth=1.2, edgecolor=f"#{GRID}", facecolor=f"#{BG_LIGHT}", zorder=1))
    ax.text(0.6, main_top - 0.28, "MAIN INTEGRATION PROCESS", fontsize=8, fontweight='bold',
            color=f"#{NAVY}", va='center', zorder=2, fontfamily='DejaVu Sans')

    # Exception panel
    err_top, err_bottom = 1.75, 0.62
    ax.add_patch(patches.FancyBboxPatch((0.35, err_bottom), fig_w - 0.7, err_top - err_bottom,
                 boxstyle="round,pad=0.02,rounding_size=0.06",
                 linewidth=1.2, edgecolor="#F2C4BC", facecolor=f"#{BG_ERR}", zorder=1))
    ax.text(0.6, err_top - 0.24, "LOCAL EXCEPTION SUBPROCESS", fontsize=8, fontweight='bold',
            color=f"#{RED}", va='center', zorder=2, fontfamily='DejaVu Sans')

    # Calculate X positions for Main Pipeline
    xs = [0.35 + gap * 0.75 + i * gap for i in range(total)]
    span = xs[-1] - xs[0] if total > 0 else 0
    offset = (fig_w - 0.7 - span) / 2 + 0.35 - xs[0] if total > 0 else 0
    xs = [x + offset for x in xs]

    # Process and Number Main Pipeline Nodes
    raw_nodes = []
    step_counter = 1
    for i, node in enumerate(diagram_nodes):
        ntype = node.get("node_type", "Processing")
        if ntype in ["Sender", "Receiver"]:
            color = GREEN
            num = None
        elif ntype == "End":
            color = SLATE
            num = None
        else:
            color = TYPE_COLORS.get(node.get("subtype", ""), NAVY_SOFT)
            num = step_counter
            step_counter += 1
            
        raw_nodes.append((node.get("name", ""), node.get("subtype", ""), xs[i], color, num))

    wrapped_nodes = []
    for label, styp, x, color, num in raw_nodes:
        label_lines = _wrap_label(label, LABEL_WRAP_CHARS, max_lines=4)
        subtype_lines = _wrap_label(styp, SUBTYPE_WRAP_CHARS, max_lines=2)
        wrapped_nodes.append((label_lines, subtype_lines, x, color, num))

    max_label_lines = max((len(n[0]) for n in wrapped_nodes), default=1)
    max_subtype_lines = max((len(n[1]) for n in wrapped_nodes), default=1)
    node_h = 0.5 + max_label_lines * 0.19 + max_subtype_lines * 0.06

    y_main = main_bottom + (main_top - main_bottom) * 0.42

    for label_lines, subtype_lines, x, color, num in wrapped_nodes:
        _rounded_node(ax, x, y_main, node_w, node_h, label_lines, subtype_lines, color, number=num)

    for i in range(len(wrapped_nodes) - 1):
        _arrow(ax, wrapped_nodes[i][2] + node_w / 2, y_main, wrapped_nodes[i + 1][2] - node_w / 2, y_main, NAVY)

    # ---------------------------------------------
    # EXCEPTION NODES (Drawn sequentially, no vertical drop arrow)
    # ---------------------------------------------
    exc_nodes = iflow_data.get("exception_nodes")
    if not exc_nodes or not isinstance(exc_nodes, list):
        exc_nodes = [{"name": "Catch Exception", "subtype": "Exception Handling"}, 
                     {"name": "Log Error", "subtype": "Exception Handling"}]
        
    e_total = len(exc_nodes)
    e_xs = [0.35 + gap * 0.75 + i * gap for i in range(e_total)]
    e_span = e_xs[-1] - e_xs[0] if e_total > 0 else 0
    e_offset = (fig_w - 0.7 - e_span) / 2 + 0.35 - e_xs[0] if e_total > 0 else 0
    e_xs = [x + e_offset for x in e_xs]

    y_err = (err_top - 0.5 + err_bottom) / 2 - 0.02
    exc_node_h = 0.5 + 3 * 0.19 

    wrapped_exc_nodes = []
    for i, enode in enumerate(exc_nodes):
        # Handle both LLM Dictionary format and original XML Parser Tuple format
        if isinstance(enode, dict):
            e_name = enode.get("name", "Exception Step")
            e_type = enode.get("subtype", "Exception Handling")
        else:
            e_name = enode[0] if len(enode) > 0 else "Exception Step"
            e_type = enode[1] if len(enode) > 1 else "Exception Handling"
            
        name_lines = _wrap_label(e_name, LABEL_WRAP_CHARS, max_lines=4)
        sub_lines = _wrap_label(e_type, SUBTYPE_WRAP_CHARS, max_lines=1)
        wrapped_exc_nodes.append((name_lines, sub_lines, e_xs[i]))

    for name_lines, sub_lines, ex_x in wrapped_exc_nodes:
        _rounded_node(ax, ex_x, y_err, node_w, exc_node_h, name_lines, sub_lines, RED)

    for i in range(len(wrapped_exc_nodes) - 1):
        _arrow(ax, wrapped_exc_nodes[i][2] + node_w / 2, y_err, wrapped_exc_nodes[i + 1][2] - node_w / 2, y_err, RED, lw=1.4)

    # Legend
    legend_items = [("Sender / Receiver", GREEN), ("Groovy Script", ORANGE),
                     ("Data Mapping", PURPLE), ("Processing Step", NAVY_SOFT),
                     ("Exception Handling", RED)]
    lx = 0.6
    for text, color in legend_items:
        ax.add_patch(patches.FancyBboxPatch((lx, 0.12), 0.22, 0.13,
                     boxstyle="round,pad=0.01,rounding_size=0.03",
                     linewidth=0, facecolor=f"#{color}", zorder=2))
        ax.text(lx + 0.3, 0.185, text, fontsize=6.6, color=f"#{SLATE}", va='center', fontfamily='DejaVu Sans')
        lx += 0.3 + len(text) * 0.072 + 0.35

    plt.savefig(output_img_path, bbox_inches='tight', facecolor=f"#{WHITE}")
    plt.close()
    return output_img_path

#===========================================================================
# WORD XML / STYLING HELPERS
# ===========================================================================

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hexcolor}"/>'))


def set_cell_borders(cell, color=GRID, sz=4, sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=90, bottom=90, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/></w:tcMar>'))


def vcenter(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_header_row(row, bg=NAVY):
    for cell in row.cells:
        set_cell_bg(cell, bg)
        set_cell_borders(cell, color=bg)
        set_cell_margins(cell, top=110, bottom=110)
        vcenter(cell)
        p = cell.paragraphs[0]
        if not p.runs:
            p.add_run("")
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb(WHITE)
            run.font.name = FONT


def zebra_rows(table, start=1, label_col_bold=True):
    for i, row in enumerate(table.rows[start:]):
        bg = WHITE if i % 2 == 0 else BG_LIGHT
        for ci, cell in enumerate(row.cells):
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            set_cell_margins(cell)
            vcenter(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = FONT
                    if label_col_bold and ci == 0:
                        run.font.bold = True
                        run.font.color.rgb = rgb(SLATE)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style=doc.styles['Heading 1'])
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), ORANGE)
    pbdr.append(bottom)
    pPr.append(pbdr)
    r_num = p.add_run(f"{number}  ")
    r_num.font.color.rgb = rgb(ORANGE)
    r_num.font.bold = True
    r_num.font.size = Pt(14)
    r_num.font.name = FONT
    r_title = p.add_run(title.upper())
    r_title.font.color.rgb = rgb(NAVY)
    r_title.font.bold = True
    r_title.font.size = Pt(14)
    r_title.font.name = FONT
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph(style=doc.styles['Heading 2'])
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.color.rgb = rgb(NAVY_SOFT)
    r.font.bold = True
    r.font.size = Pt(11.5)
    r.font.name = FONT
    return p


def body_text(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = FONT
    r.font.color.rgb = rgb(TEXT_BODY)
    return p


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = FONT
        r.font.color.rgb = rgb(TEXT_BODY)
    return


CONTENT_WIDTH_IN = 6.9  # page width minus L/R margins - matches header/footer tables


def set_col_widths(table, widths_in):
    """Force fixed column widths so Word can't auto-shrink a column to make
    room for a long unbroken value in the next one."""
    table.autofit = False
    for ci, w in enumerate(widths_in):
        table.columns[ci].width = Inches(w)
    for row in table.rows:
        for ci, w in enumerate(widths_in):
            if ci < len(row.cells):
                row.cells[ci].width = Inches(w)


def default_col_widths(ncols):
    if ncols == 2:
        return [2.3, CONTENT_WIDTH_IN - 2.3]
    w = CONTENT_WIDTH_IN / ncols
    return [w] * ncols


def new_kv_table(doc, rows, ncols=2, col_widths=None):
    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for values in rows:
        cells = table.add_row().cells
        for i, val in enumerate(values):
            cells[i].text = str(val)
    zebra_rows(table, start=0)
    set_col_widths(table, col_widths or default_col_widths(ncols))
    return table


def new_headed_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
    style_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, val in enumerate(values):
            cells[i].text = str(val)
    zebra_rows(table, start=1, label_col_bold=False)
    set_col_widths(table, col_widths or default_col_widths(len(headers)))
    return table


def new_security_table(doc, rows):
    widths = [2.4, CONTENT_WIDTH_IN - 2.4]
    return new_headed_table(doc, ["Security Material Type", "Alias / Name"], rows, col_widths=widths)


def add_field(run, field_str, default_text=""):
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_str} '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    if default_text:
        t = OxmlElement('w:t'); t.text = default_text
        run._r.append(t)
    run._r.append(fld_end)


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    add_field(r, 'TOC \\o "1-2" \\h \\z \\u',
              default_text="Right-click here and choose \u201cUpdate Field\u201d to build the table of contents.")


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    el = OxmlElement('w:updateFields')
    el.set(qn('w:val'), 'true')
    settings.append(el)


# ===========================================================================
# DOCUMENT BUILDER
# ===========================================================================

def build_visteon_ts_docx(iflow_data, logo_path=PRIMARY_LOGO_PATH, output_dir="."):
    doc = Document()
    iflow_data = apply_user_metadata(iflow_data)
    iflow_name = (iflow_data.get("iflow_name") or "SAP_CPI_IFLOW").strip()
    safe_iflow_name = sanitize_output_name(iflow_name)
    output_dir = os.path.abspath(output_dir or ".")
    os.makedirs(output_dir, exist_ok=True)
    output_doc_path = os.path.join(output_dir, f"UID_U057_{safe_iflow_name}.docx")
    resolved_logo_path = resolve_asset_path(logo_path)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        header = section.header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.9))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = True

        cell_left = header_table.rows[0].cells[0]
        set_cell_borders(cell_left, color=NAVY, sz=6, sides=("bottom",))
        p_logo = cell_left.paragraphs[0]
        if resolved_logo_path and os.path.exists(resolved_logo_path):
            p_logo.add_run().add_picture(resolved_logo_path, width=Inches(1.6))
        else:
            r_fb = p_logo.add_run("VISTEON")
            r_fb.font.bold = True
            r_fb.font.size = Pt(14)
            r_fb.font.name = FONT
            r_fb.font.color.rgb = rgb(NAVY)

        cell_right = header_table.rows[0].cells[1]
        set_cell_borders(cell_right, color=NAVY, sz=6, sides=("bottom",))
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_hdr = p_right.add_run(f"{iflow_name[:44]}")
        r_hdr.font.size = Pt(9)
        r_hdr.font.bold = True
        r_hdr.font.name = FONT
        r_hdr.font.color.rgb = rgb(NAVY)
        p_right.add_run().add_break()
        r_hdr2 = p_right.add_run("Technical Specification")
        r_hdr2.font.size = Pt(7.5)
        r_hdr2.font.name = FONT
        r_hdr2.font.color.rgb = rgb(SLATE)

        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.9))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        f_left_cell = footer_table.rows[0].cells[0]
        set_cell_borders(f_left_cell, color=GRID, sz=6, sides=("top",))
        f_left = f_left_cell.paragraphs[0]
        r_fl = f_left.add_run("Visteon Internal & Confidential  |  SAP Integration Suite")
        r_fl.font.size = Pt(7.5)
        r_fl.font.name = FONT
        r_fl.font.color.rgb = rgb("808080")

        f_right_cell = footer_table.rows[0].cells[1]
        set_cell_borders(f_right_cell, color=GRID, sz=6, sides=("top",))
        f_right = f_right_cell.paragraphs[0]
        f_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_fr1 = f_right.add_run("Page ")
        r_fr1.font.size = Pt(7.5); r_fr1.font.name = FONT; r_fr1.font.color.rgb = rgb("808080")
        add_field(r_fr1, "PAGE", "1")
        r_fr2 = f_right.add_run(" of ")
        r_fr2.font.size = Pt(7.5); r_fr2.font.name = FONT; r_fr2.font.color.rgb = rgb("808080")
        add_field(r_fr2, "NUMPAGES", "1")

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(TEXT_BODY)

    for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        hs = doc.styles[style_name]
        hs.font.name = FONT
        hs.font.color.rgb = rgb(NAVY)

    # ---------------- COVER PAGE ----------------
    cov_top = doc.add_paragraph()
    cov_top.paragraph_format.space_before = Pt(130)
    if resolved_logo_path and os.path.exists(resolved_logo_path):
        run_logo = cov_top.add_run()
        run_logo.add_picture(resolved_logo_path, width=Inches(2.1))
    else:
        r_logo_fb = cov_top.add_run("VISTEON")
        r_logo_fb.font.bold = True
        r_logo_fb.font.size = Pt(22)
        r_logo_fb.font.name = FONT
        r_logo_fb.font.color.rgb = rgb(NAVY)

    cov_kicker = doc.add_paragraph()
    cov_kicker.paragraph_format.space_before = Pt(36)
    r_kicker = cov_kicker.add_run("SAP INTEGRATION SUITE  |  TECHNICAL SPECIFICATION")
    r_kicker.font.size = Pt(11)
    r_kicker.font.bold = True
    r_kicker.font.name = FONT
    r_kicker.font.color.rgb = rgb(ORANGE)

    cov_title = doc.add_paragraph()
    cov_title.paragraph_format.space_before = Pt(6)
    cov_title.paragraph_format.space_after = Pt(4)
    r_title = cov_title.add_run(iflow_name)
    r_title.font.size = Pt(26)
    r_title.font.bold = True
    r_title.font.name = FONT
    r_title.font.color.rgb = rgb(NAVY)

    cov_sub = doc.add_paragraph()
    r_sub = cov_sub.add_run(f"{iflow_data['sender_system']}  \u2192  {iflow_data['receiver_system']}")
    r_sub.font.size = Pt(13)
    r_sub.font.name = FONT
    r_sub.font.color.rgb = rgb(SLATE)

    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    cov_meta = new_kv_table(doc, [
        ("Integration Package", iflow_data['package_name']),
        ("Document Version", iflow_data['doc_version']),
        ("Prepared By", iflow_data['prepared_by']),
        ("Reviewed By", iflow_data['reviewed_by']),
        ("Approved By", iflow_data['approved_by']),
        ("Effective Date", iflow_data['effective_date']),
        ("Classification", "Visteon Internal & Confidential"),
    ])
    for row in cov_meta.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.5)

    doc.add_page_break()

    # ---------------- TABLE OF CONTENTS ----------------
    toc_head = doc.add_paragraph()
    r_toc = toc_head.add_run("TABLE OF CONTENTS")
    r_toc.font.size = Pt(16)
    r_toc.font.bold = True
    r_toc.font.name = FONT
    r_toc.font.color.rgb = rgb(NAVY)
    add_toc(doc)
    doc.add_page_break()

    # ---------------- REVISION HISTORY ----------------
    add_section_heading(doc, "", "Revision History")
    rev_headers = ["Version", "Effective Date", "Brief Description", "Change Ref", "Affected Section",
                   "Prepared By", "Reviewed By", "Approved By"]
    new_headed_table(doc, rev_headers,
                      [[iflow_data['doc_version'], iflow_data['effective_date'],
                        iflow_data['description'],
                        "-", "All Sections",
                        iflow_data['prepared_by'],
                        iflow_data['reviewed_by'],
                        iflow_data['approved_by']]])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 1. OVERVIEW ----------------
    add_section_heading(doc, "1.", "Overview")
    overview_text = iflow_data.get('business_overview') or iflow_data.get('interface_description') or (
        f"This document outlines the end-to-end integration setup where data flows from {iflow_data['sender_system']} "
        f"to {iflow_data['receiver_system']} using SAP Cloud Integration (SAP CPI)."
    )
    body_text(doc, overview_text)

    # ---------------- 2. INTERFACE FLOW ----------------
    add_section_heading(doc, "2.", "Interface Flow")


    # Determine schedule triggers based on Execution Mode, not just Adapter type
    is_scheduled = "Scheduled" in iflow_data.get('execution_mode', '')
    trigger_event = "Timer Schedule / Polling" if is_scheduled else "Inbound Request"
    
    # Use LLM extracted frequency, or fallback intelligently
    freq_run = iflow_data.get('frequency')
    if not freq_run or "request" in freq_run.lower() and is_scheduled:
        freq_run = "Configured Cron Schedule" if is_scheduled else "Immediate upon request"

    flow_info = [
        ("Source System Name", iflow_data['sender_system']),
        ("Target System Name", iflow_data['receiver_system']),
        ("Direction of Interface", iflow_data.get('direction', 'Inbound')),
        ("Execution Mode", iflow_data['execution_mode']),
        ("Synchronous / Asynchronous", iflow_data['synchronous_asynchronous']),
        ("Source Transport Protocol", iflow_data['sender_adapter_type']),
        ("Target Transport Protocol", iflow_data['receiver_adapter_type']),
        ("Triggering Events", trigger_event),
        ("Frequency of Run", freq_run),
        ("Volume", "As per operational schedule"),
        ("Acknowledgement", "No"),
    ]
    new_kv_table(doc, flow_info)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 3. ASSUMPTIONS ----------------
    add_section_heading(doc, "3.", "Assumptions")
    add_bullets(doc, [
        f"Data flows seamlessly from {iflow_data['sender_system']} to "
        f"{iflow_data['receiver_system']} according to configured message transformations."
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 4. INTERFACE DESCRIPTION ----------------
    add_section_heading(doc, "4.", "Interface Description")
    description_text = iflow_data.get('interface_description') or iflow_data.get('description') or (
        f"This integration facilitates the transfer of operational data from {iflow_data['sender_system']} to "
        f"{iflow_data['receiver_system']} via SAP CPI."
    )
    body_text(doc, description_text)

    # ---------------- 5. TECHNICAL ARTIFACTS ----------------
    add_section_heading(doc, "5.", "Technical Artifacts")
    groovy_txt = "\n".join(f"{i+1}. {s}" for i, s in enumerate(iflow_data["groovy_scripts"])) or "None"
    xslt_txt = "\n".join(f"{i+1}. {s}" for i, s in enumerate(iflow_data["xslt_scripts"])) or ""
    mappings_txt = "\n".join(f"{i+1}. {s}" for i, s in enumerate(iflow_data["mappings"])) or ""
    mapping_combined = "\n".join(t for t in [xslt_txt, mappings_txt] if t) or "None"

    new_kv_table(doc, [
        ("Integration Package", iflow_data['package_name']),
        ("Integration iFlow", iflow_name),
        ("Sender System", iflow_data['sender_system']),
        ("Receiver System", iflow_data['receiver_system']),
        ("XSLT / Message Mappings", mapping_combined),
        ("Groovy Scripts", groovy_txt),
        ("Business Identifier", "Application ID / Track ID"),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 6. HIGH LEVEL PROCESSING LOGIC ----------------
    # ---------------- 6. HIGH LEVEL PROCESSING LOGIC ----------------
    add_section_heading(doc, "6.", "High Level Processing Logic")
    
    processing_logic = iflow_data.get('processing_logic')
    
    # If the LLM returned a structured list, print each point as a bullet
    if isinstance(processing_logic, list) and processing_logic:
        for step_text in processing_logic:
            # Clean up any stray brackets/quotes just in case
            clean_text = re.sub(r"^['\"]+|['\"]+$", "", step_text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(clean_text)
            r.font.size = Pt(10)
            r.font.name = FONT
            r.font.color.rgb = rgb(TEXT_BODY)
    else:
        # Fallback if standard string is returned
        processing_text = processing_logic if isinstance(processing_logic, str) else (
            f"Data from {iflow_data['sender_system']} is received and processed through {iflow_name} in SAP CPI. "
            f"The flow executes custom logic, transforms the payload structure, and delivers messages to "
            f"{iflow_data['receiver_system']}."
        )
        body_text(doc, processing_text)

    # ---------------- 7. PROCESS FLOW DIAGRAM ----------------
    add_section_heading(doc, "7.", "Process Flow Diagram")
    diagram_path = os.path.join(output_dir, f"diag_{safe_iflow_name}.png")
    img_path = draw_enterprise_diagram(iflow_data, output_img_path=diagram_path)
    doc.add_picture(img_path, width=Inches(6.6))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 8. SECURITY MATERIAL ----------------
    add_section_heading(doc, "8.", "Security Material & Credentials")
    security_rows = [[t, a] for t, a in iflow_data["security_materials"]]
    if not security_rows:
        security_rows = [["None detected", "-"]]
    new_security_table(doc, security_rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 9. CONFIGURATION PARAMETERS ----------------
    add_section_heading(doc, "9.", "Configuration Parameters (Dev)")

    add_sub_heading(doc, f"9.1  Sender Adapter / Protocol ({iflow_data['sender_adapter_type']})")
    sender_p = iflow_data["sender_params"] or [("Sender", iflow_data['sender_system']),
                                                ("Adapter Type", iflow_data['sender_adapter_type'])]
    new_headed_table(doc, ["Parameter Key", "Configured Value"], sender_p)

    add_sub_heading(doc, f"9.2  Receiver Adapter / Protocol ({iflow_data['receiver_adapter_type']})")
    receiver_p = iflow_data["receiver_params"] or [("Receiver", iflow_data['receiver_system']),
                                                     ("Adapter Type", iflow_data['receiver_adapter_type'])]
    new_headed_table(doc, ["Parameter Key", "Configured Value"], receiver_p)

    add_sub_heading(doc, "9.3  Externalized Properties")
    props = iflow_data["properties"] or {"Receiver_System": iflow_data['receiver_system']}
    new_headed_table(doc, ["Property Key", "Configured Expression / Value"], list(props.items()))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 10. DATA MAPPING ----------------
    add_section_heading(doc, "10.", "Data Mapping Matrix")
    new_headed_table(doc, ["Source Field Name", "Target Field Name", "Mapping Logic", "Comments"],
                      [["SourcePayload", "TargetPayload", "Transformation via Mapping Artifacts",
                        "Automated Field Mapping"]])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 11. INTERFACE ALERT NOTIFICATION ----------------
    add_section_heading(doc, "11.", "Interface Alert Notification")
    new_headed_table(doc, ["Alert Parameter", "Configuration Details"],
                      [["Error Handler Subprocess", "Global Exception Handler Process"],
                       ["Alert Channel", "CPI Operational Monitoring Dashboard / Email"]])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 12. ERROR CONDITIONS ----------------
    add_section_heading(doc, "12.", "Error Conditions, Message Retry & Reprocessing")
    new_kv_table(doc, [
        ("Message Retry Subprocess", "Standard SAP CPI Automatic Retry"),
        ("Max Interval", "Exponential Backoff Configuration"),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 13. UNIT TEST DOCUMENTS ----------------
    add_section_heading(doc, "13.", "Unit Test Documents")
    body_text(doc, f"UT_DOCUMENT_U057_{iflow_name}.docx")

    # ---------------- 14. INTERFACE MONITORING ----------------
    add_section_heading(doc, "14.", "Interface Monitoring Procedures")
    body_text(doc, "Monitor message execution status, MPL logs, and Exception Handler queues in the "
                    "SAP Integration Suite Monitoring dashboard.")

    # ---------------- 15. REFERENCES ----------------
    add_section_heading(doc, "15.", "Project References")
    new_headed_table(doc, ["Sr. No.", "Document Name", "Reference Link"], [
        ["", "", ""],
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------- 16. COMMENTS & APPROVALS ----------------
    add_section_heading(doc, "16.", "Comments & Approvals")
    new_headed_table(doc, ["Date", "Raised By", "Comment", "Status"], [
        ["", "", "", ""],
        ["", "", "", ""],
        ["", "", "", ""],
    ])

    set_update_fields_on_open(doc)
    doc.save(output_doc_path)
    print(f"Generated: {output_doc_path}")
    return output_doc_path



# ===========================================================================
# BATCH RUNNER
# ===========================================================================

def process_all_cpi_zips(directory_path):
    if not os.path.exists(directory_path):
        print(f"Target directory '{directory_path}' does not exist.")
        return

    zip_files = glob.glob(os.path.join(directory_path, "*.zip"))
    if not zip_files:
        print(f"No .zip files found in {directory_path}")
        return

    print(f"Found {len(zip_files)} CPI ZIP file(s). Generating Technical Specification documents...")
    for zip_file in zip_files:
        parsed_data = parse_cpi_iflow_zip(zip_file)
        build_visteon_ts_docx(parsed_data)

if __name__ == "__main__":
    process_all_cpi_zips(ZIP_DIRECTORY_PATH)